#This is a Python3 script to for checking the word choice of documents of text.
#Every word in a (*.txt or *.doc) file is put into a list of tuples (String word, int count)
#Every word in a paragraph is put into a list of tuples (String word, int count)
#Every word in a sentence is put into a list of tuples (String word, int count)
#Every word is then scored accordingly in a dictionary:
#     for every word a sentence contains more than once, it receives a score of +3
#     for every word a paragraph contains more than once, it recieves a score of +2
#     for every word that is used in a file, it receive a score of its total word count.
#The user is then presented these words in order of score, and is allowed to cycle back and forth through the list of words
#The user may choose to exit the application by entering "0" at any time.

############################################################
####################    IMPORTS         ####################
############################################################
import os#necessary to get the directory of the file in question (and the directory of this program).
from docx import Document#necessary to read MS docx files!  WIthout it, limited to only reading text files ans ASCII characters!

############################################################
####################    FUNCTIONS       ####################
############################################################
#prompts the user for a file name and extension.  Once a valid file is input, the parent directory, name, and extension are returned as a tuplefrom docx import Document # necessary
#@return (String, String, String) returns a tuple of three strings, the parent directory, the file name, and the file extension.
def getFileAttributes():
    filePath = os.path.dirname(os.path.realpath(__file__))
    fullFileName = input("Enter the name and extension of the file (ex. Some File.docx).\nIMPORTANT:  It must be in the same directory as this program!:\n\n    >>>" + filePath + "\\")
    fileName, fileExtension = "", ""
    beforeExtension = True
    if (fullFileName.count(".") != 1):
        print("ERROR, PROBLEM WITH FILE EXTENSION (check period!).\n\n\n")
        return getFileAttributes()
    for char in fullFileName:
        if char != "." and beforeExtension:
            fileName += char
        else:
            beforeExtension = False
            fileExtension += char
    if fileExists(filePath, fileName, fileExtension) == False:
        print("ERROR, THAT FILE COULD NOT BE FOUND.\n\n\n")
        return getFileAttributes()
    return (filePath, fileName, fileExtension)

#checks if a specified file exists.  Returns true if the file is found, and false if the file is not
#@param String filePath is the parent directory of the file
#@param String fileName is the name of the file (does not include the parent directory or extension!
#@param String fileExtension is the extension of the file (ex. ".txt")
#@return bool return true if the file if found (ie exists), and false otherwise.
def fileExists(filePath, fileName, fileExtension, action = 'r'):
    try:
        f = open(filePath + "\\" + fileName + fileExtension, action)
        f.close()
        return True;
    except:
        try:
            doc = Document(open(filePath + "\\" + fileName + fileExtension, action + 'b'))
            doc.close()
            return True
        except FileNotFoundError:
            return False
        return False
    return True

#Appends a line to a specified file (assumed in working directory), and also adds additional functionality for the addiion.  It is assumed, however, that the appending is standard
#@String s is the string representing the text to be added to the file
#@String fileName is the name of the dile to which the text will be appended.  It is assumed to be in the working directory.
#@bool allowDuplicates is a boolean value which is true by default.  If specified to false, the string being added must be an original line in the file or else it will not be added
#@bool alphabetizeFile is a boolean value which is false by default.  If specified to be true, the entire file will be alphabetized by line.  the current file will then be overwritten with the newly alphabetized version.
#@return void
def addNewLine(s, fileName, allowDuplicates = True, alphabetizeFile = False):#return none;  appends string s to a new line in file fileName and then rewrites the file so it is in alpabetical order.  CURRENTLY NOT EFFIECIENT FOR LARGE FILE SIZES.  Assumes that s is a verified string.
    if isKnown(s, fileName):#checks if the element to be added is already in the file.  Also checks that there is a file to add to!  (if not, creates the file!)
        if allowDuplicates == False:#from encasing (previous) if statement, know that there is a dupliate in there! So no need to add, just return!
            return#no need to add anthing
        f = open(fileName, 'a')#append this file
        f.write(s + "\n")#add the element to the file before sorting it.
        f.close()
    else: #file (or string in that file) does not exist!  no need to D2!
        f = open(fileName, 'a')#append this file
        f.write(s + "\n")#add the element to the file before sorting it.
        print("Appended " + s + " to " + fileName)
        f.close()
    if alphabetizeFile == True:
        #need to enter the name so it is in alphabetical order!
        f = open(fileName, 'rb')
        tmpList = []
        for line in sorted(f):
            tmpList.append(line)
            print("#####" + line)
        f.close()
        f = open(fileName, 'w')
        for element in tmpList:
            f.write(element)
        f.close()

#safely creates a file of the specified name (assumed to be in the working directory!)
#@String fileName is the name of the file to be created.  It will be created by in the current working directory
#@return void
def createFile(fileName):
    f = open(fileName, 'a')
    f.close

#returns all the text of a file as a single string.
#@String directory is the parent directory of the file specified in fileName
#@return void
def fileToString(directory, fileName, extension):
    #######################
    ### LOCAL FUNCTIONS ###
    #######################
    def isAscii(s):
        try:
            s.encode('ascii')
        except UnicodeEncodeError:
            return False
        else:
            return True

    ###########################
    ###        MAIN         ###
    ###########################
    eraseTmpFile = False
    TMP_FILENAME = directory + "tmpFile" + '.txt'
    if (extension == ".txt"):
        f = open(directory + "\\" + fileName + extension, 'r')
    elif (extension == ".docx"):
        print ("\n Looking for the following doc: (" + directory + "\\) " + fileName + extension + "\n")
        doc = Document(open(directory + "\\" + fileName + extension, 'rb'))
        createFile(TMP_FILENAME)
        eraseTmpFile = True
        tmpf = open(TMP_FILENAME, 'a')
        for para in doc.paragraphs:
            try:
                tmpf.write(para.text)
            except UnicodeEncodeError:
                tmpf.write(" ")
            tmpf.write("\n")#turn string to txt file
        tmpf.close()
        f = open(TMP_FILENAME, 'r')
    s = "";
    for line in f:
        for char in line:
            if (isAscii(str(char))):
                s += str(char)
            else:
                s+= " "
    f.close()
    if (eraseTmpFile):
        os.remove(TMP_FILENAME)
    return s

def updateScores(l, d, bonus):
    alreadyUsedElts = [" "]
    alreadyUsedElts.clear()
    for elt in l:
        if elt in d:
            if alreadyUsedElts.count(elt) > 0:
                d[elt] += bonus
            else:
                alreadyUsedElts.append(elt)
                d[elt] += 1
        else:
            alreadyUsedElts.append(elt)
            d[elt] = 1
    return d


def delimitStringToList(s):
    FIRST_ELT = " "
    l = [FIRST_ELT]
    l.clear()
    tmpStr = ""
    for char in s:
        if isRunonPunctuation(char) or isMiscPunctuation(char) or char == " ":
            if tmpStr != "":
                l.append(tmpStr)
            tmpStr = ""
        elif char == "\n":
            if tmpStr != "":
                l.append(tmpStr)
            l.append(char)
            tmpStr = ""
        else:
            tmpStr += char
    if tmpStr != "":
        l.append(tmpStr)
    return l

def isEndingPunctuation(c):
    return c == '.' or c == '!' or c == '?' or c == ';'
def isRunonPunctuation(c):
    return c == "," or c == "\t" or c == ':'
def isMiscPunctuation(c):
    return c == "\"" or c == "(" or c == ")"
def isPunctuation(c):
    return isEndingPunctuation(c) or isRunonPunctuation(c) or isMiscPunctuation(c)

def convertToSortedList(d):
    l = [""]
    l.clear()
    for k,v in d.items():
        l.append((v, k))
    return sorted(l, key=lambda x: x[0])

def getLength(l):
    totalWords = 0
    for elt in l:
        totalWords += 1
    return totalWords


############################################################
####################        MAIN        ####################
############################################################
FIRST_ELT = " "
SENTENCE_WEIGHT = 10;
PARAGRAPH_WEIGHT = 4;
FILE_WEIGHT = 1;
f = getFileAttributes()
while(True):
    s = fileToString(f[0], f[1], f[2]).upper()
    allWords = delimitStringToList(s)#takes out all spaces, only leaving returns
    ###THere has to be a better way to initialize empty lists...
    fileWordCount = [FIRST_ELT]
    fileWordCount.clear()
    paragraphWordCount = [FIRST_ELT]
    paragraphWordCount.clear()
    sentenceWordCount = [FIRST_ELT]
    sentenceWordCount.clear()

    allWordScores = {}
    for word in allWords:
        #remove punctuation from the end of words
        if isEndingPunctuation(word[-1]) == False:
            sentenceWordCount.append(word)
        else:#move onto new sentence
            word = word[:-1]
            sentenceWordCount.append(word)#add the last word minus the punctuation
            allWordScores = updateScores(sentenceWordCount, allWordScores, SENTENCE_WEIGHT)
            sentenceWordCount.clear()

        #separate into paragraphs (the returns have already been removed and are treated as individual words!)
        if word != "\n" and word != "\r\n":
            paragraphWordCount.append(word)
        else:#move onto new paragraph
            allWordScores = updateScores(paragraphWordCount, allWordScores, PARAGRAPH_WEIGHT)
            paragraphWordCount.clear()

        #update file Count
        fileWordCount.append(word)

    allWordScores = updateScores(fileWordCount, allWordScores, FILE_WEIGHT)
    allWordScores = updateScores(paragraphWordCount, allWordScores, PARAGRAPH_WEIGHT)
    allWordScores = updateScores(sentenceWordCount, allWordScores, SENTENCE_WEIGHT)

    allWords = {}
    allWords = updateScores(fileWordCount, allWords, 1)

    print("\n+____________________________________________\n")
    print (" _Freq_\t_Total_\t_Word_")
    totalWords = getLength(convertToSortedList(allWordScores))
    count = 0
    for elt in convertToSortedList(allWordScores):
        if elt[1] != "\n":
            print (" " + str(elt[0]) + "\t" + str(allWords[elt[1]]) + "\t" + elt[1])
        else:
            print (" " + str(elt[0]) + "\t" + str(allWords[elt[1]]) + "\t" + "<CARRIAGE RETURN>")
        count += 1
        if ((totalWords - count) % 25 == 0) and (totalWords - count >= 5):
            print(" ----------- Top " + str(totalWords - count) + " -----------")

    print("Scoring complete.\nTotal Vocab used: " + str(totalWords) + " different words.\n")
    input("Hit \'enter\' to refresh the evaluation\n\n")
print("Goodbye.")
exit(0)
